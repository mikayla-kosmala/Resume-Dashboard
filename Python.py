from docx import Document
import string
import re
import pandas as pd
from datetime import datetime

translator = str.maketrans('', '', string.punctuation)
# Load your resume
doc = Document(r'C:\Users\mikay\OneDrive\Documents\Resume to Tableau\Running Resume.docx')

# Print all the paragraphs
#for para in doc.paragraphs:
#  print(para.text)
job = []
job_info = []
job_title = []
company = []
start_end_date = []
job_desc = []
job_achievements = []
found_title = 0
found_company = 0
found_date = 0
found_desc = 0
found_achievements = 0
desc_found =0
section = ''
df = pd.DataFrame(columns=["Section","Title", "Company", "Desc", "Accomplishments", "Start Date", "End Date"])
for para in doc.paragraphs:
    print(f"Style: {para.style.name} | Text: {para.text}")
    if "Heading" in para.style.name:
        section = para.text
    if section == "Experience":
        #print(para.text)
        if desc_found:
            desc = para.text
            desc_found = 0
        if "Bullet List" in para.style.name:
            # if end == " Current":
            #     end = datetime.today().strftime("%b %Y")
            new_row = {"Section": section, "Title": title, "Company": company, "Desc": desc, "Accomplishments":para.text, "Start Date":start, "End Date":end}
            # Append the row
            df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        for run in para.runs:
            if run.bold:
                text = run.text.strip()
                if text:  # Ignore empty strings
                    title = text
                    list = re.sub(r'\t+','|',para.text[len(text)+1:]).split('|')
                    company=list[0]
                    start,end=list[1].split('–')
                    desc_found = 1
                    break
    if section == "Projects":
        if desc_found:
            desc = para.text.split('|')[1:]
            start= para.text.split('|')[0]
            end = start
            if '-' in start:
                start,end= start.split('-')
            desc_found = 0
        if "Bullet List" in para.style.name:
            new_row = {"Section": section, "Title": title, "Company": company, "Desc": desc, "Accomplishments":para.text, "Start Date":start, "End Date":end}
            # Append the row
            df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        for run in para.runs:
            if run.bold:
                text = run.text.strip()
                if text:  # Ignore empty strings
                    title = text
                    list = re.sub(r'\t+','|',para.text[len(text)+2:]).split('|')
                    company=list[0]
                    print(list)
                    desc_found = 1
                    break
    if section == "Education":
        for run in para.runs:
            print(f"Text: '{run.text}' | Bold: {run.bold} | Italic: {run.italic}")
            text = run.text.strip()
            if run.bold and found_company==0:
                if text:  # Ignore empty strings
                    company = text
                    print('Company: ', company)
                    found_company = 1
            if run.italic and found_desc==0:
                text = run.text.strip()
                if text:
                    desc = text
                    print('Desc: ', desc)
                    found_desc = 1
            if  not run.bold and not run.italic and not ("Heading" in para.style.name) and text and found_company and found_desc: 
                print(text)
                list = re.sub(r'\t+','|',para.text[len(text)+2:]).split('|')
                print(list)
                start,end=list[1].split('–')
                new_row = {"Section": section, "Title": 'Student', "Company": company[:-1], "Desc": desc, "Accomplishments":'', "Start Date":start, "End Date":end}
                # Append the row
                df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
                found_company = 0
                found_desc = 0
                

print(df.head(70))

df.to_excel(r"C:\Users\mikay\OneDrive\Documents\Resume to Tableau\resume_data.xlsx", index=False)
#Second Attempt
# for para in doc.paragraphs:
#     print(f"Style: {para.style.name} | Text: {para.text}")
#     #print(para.text)
#     if "Bullet List" in para.style.name:
        
#     for run in para.runs:
#         if run.bold:
#             text = run.text.strip()
#             if text:  # Ignore empty strings
#                 job_title.append(text)
                
#                 list = re.sub(r'\t+','|',para.text[len(text)+2:]).split('|')
#                 print(list)
#                 company.append(list[0])
#                 start_end_date.append(list[1])
#                 break
       



#print(job_title)
#print (company)

#FIRST ATTEMPT
# for para in doc.paragraphs:
#     print("•", para.text)
#     for run in para.runs:
#         text = run.text.strip()
#         #print(text)
#         if 'List Bullet' in para.style.name:
#             print("•", para.text.strip())
#         if found_date:
#             if text:
#                 job_desc.append(text)
#                 found_date = 0
#         if found_company:
#             if text:
#                 start_end_date.append(text)
#                 found_company = 0
#                 found_date = 1
#         if found_title:
#             if text:
#                 company.append(text.translate(translator)[1:])
#             found_title = 0
#             found_company = 1
#         if run.bold:
#             if text:  # Ignore empty strings
#                 job_title.append(text)
#                 found_title = 1


#print(job_desc)   


    
# for para in doc.paragraphs:
#     if "Heading" in para.style.name:
#         print(">> Section:", para.text)


        
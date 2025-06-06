from docx import Document
import string
import re
import pandas as pd
from datetime import datetime
from lxml import etree
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
# Access the relationships

translator = str.maketrans('', '', string.punctuation)
# Load your resume
doc = Document(r'C:\Users\mikay\OneDrive\Documents\Resume to Tableau\Running Resume.docx')
rels = doc.part.rels

def get_hyperlinks_from_para(para):
    hyperlinks = []

    # Loop through all hyperlink elements in the paragraph
    for hyperlink in para._element.findall(".//w:hyperlink", para._element.nsmap):
        # Get the r:id (relationship ID) that links to the actual URL
        r_id = hyperlink.get(qn('r:id'))
        if r_id:
            # Get the actual target (URL) from the paragraph's part
            url = para.part.rels[r_id].target_ref

            # Try to extract the visible link text
            texts = [node.text for node in hyperlink.findall(".//w:t", para._element.nsmap) if node.text]
            link_text = "".join(texts)

            hyperlinks.append((link_text, url))

    return hyperlinks
# Print all the paragraphs
#for para in doc.paragraphs:
#  print(para.text)
job = []
job_info = []
job_title = []
company = []
start_end_date = []
job_desc = []
job_achievements = []
found_title = 0
found_company = 0
found_date = 0
found_desc = 0
found_achievements = 0
desc_found =0
section = ''
resume_df = pd.DataFrame(columns=["Section","Title", "Company", "Desc", "Accomplishments", "Start Date", "End Date"])
skills_df = pd.DataFrame(columns=["Skill","Level"])
personal_df = pd.DataFrame(columns=["Section", "Information","Interest Level","Links","Path"])
new_rows = [{"Section":'Personal', "Information":'Resume',"Interest Level":'',"Links":'https://docs.google.com/document/d/1ehoNrqLzMcSuB2BzAyix7t51HYPdDZ7a/edit?usp=sharing&ouid=100832385938879723557&rtpof=true&sd=true','Path':""}]
personal_df = pd.concat([personal_df, pd.DataFrame(new_rows)], ignore_index=True)
for para in doc.paragraphs:
    
    if "Heading" in para.style.name:
        section = para.text
    if section == '' and 'Mikayla.Kosmala' in para.text:
        print(get_hyperlinks_from_para(para))
        for label, url in get_hyperlinks_from_para(para): 
            if ':' in label:
                information, interest_level = label.split(':')
            else:
                information = 'Email'
                interest_level = label
            new_rows = [{"Section":'Personal', "Information":information,"Interest Level":interest_level,"Links":url,'Path':""}]
            personal_df = pd.concat([personal_df, pd.DataFrame(new_rows)], ignore_index=True)
    if section == "Experience":
        #print(para.text)
        if desc_found:
            desc = para.text
            desc_found = 0
        if "Bullet List" in para.style.name:
            # if end == " Current":
            #     end = datetime.today().strftime("%b %Y")
            new_row = {"Section": section, "Title": title, "Company": company, "Desc": desc, "Accomplishments":para.text, "Start Date":start, "End Date":end}
            # Append the row
            resume_df = pd.concat([resume_df, pd.DataFrame([new_row])], ignore_index=True)
        for run in para.runs:
            if run.bold:
                text = run.text.strip()
                if text:  # Ignore empty strings
                    title = text
                    list = re.sub(r'\t+','|',para.text[len(text)+1:]).split('|')
                    company=list[0]
                    start,end=list[1].split('–')
                    desc_found = 1
                    break
    if section == "Projects":
        print(f"Style: {para.style.name} | Text: {para.text}")
        if desc_found:
            start, end = para.text.split('|')[0].split(' – ')
            desc, url = get_hyperlinks_from_para(para)[0]
            desc_found = 0
        if "Bullet List" in para.style.name:
            new_row = {"Section": section, "Title": title, "Company": company, "Desc": desc, "Accomplishments":para.text, "Start Date":start, "End Date":end,"Link":url}
            print(new_row)
            # Append the row
            resume_df = pd.concat([resume_df, pd.DataFrame([new_row])], ignore_index=True)
        for run in para.runs:
            if run.bold:
                text = run.text.strip()
                if text:  # Ignore empty strings
                    title = text
                    list = re.sub(r'\t+','|',para.text[len(text)+2:]).split('|')
                    company=list[0]
                    #print(list)
                    desc_found = 1
                    break
    if section == "Education":
        for run in para.runs:
            #print(f"Text: '{run.text}' | Bold: {run.bold} | Italic: {run.italic}")
            text = run.text.strip()
            if run.bold and found_company==0:
                if text:  # Ignore empty strings
                    company = text
                    #print('Company: ', company)
                    found_company = 1
            if run.italic and found_desc==0:
                text = run.text.strip()
                if text:
                    desc = text
                    #print('Desc: ', desc)
                    found_desc = 1
            if  not run.bold and not run.italic and not ("Heading" in para.style.name) and text and found_company and found_desc: 
                #print(text)
                list = re.sub(r'\t+','|',para.text[len(text)+2:]).split('|')
                #print(list)
                start,end=list[1].split('–')
                new_row = {"Section": section, "Title": 'Student', "Company": company[:-1], "Desc": desc, "Accomplishments":'', "Start Date":start, "End Date":end}
                # Append the row
                resume_df = pd.concat([resume_df, pd.DataFrame([new_row])], ignore_index=True)
                found_company = 0
                found_desc = 0
    if section == "Skills" and para.text != 'Skills':
            list = para.text.split('\n')
            for item in list:
                print('item ',item)
                if para.text:
                    group = item.split(':')
                    level = group[0]
                    skills = group[1].split('·')
                    #print(skills)
                    for skill in skills:
                        #print(skill)
                        new_row = {"Skill": skill, "Level": level}
                        skills_df = pd.concat([skills_df, pd.DataFrame([new_row])], ignore_index=True)
    if section == "Interests"and para.text != 'Interests':
        if para.text:
            #link = para.text.split(' ')[-1]
            interest_level=para.text.split(' ')[-1]
            information=' '.join(para.text.split(' ')[:-1])
            new_rows = [{"Section":section, "Information":information,"Interest Level":interest_level,"Links":"",'Path':1},{"Section":section, "Information":information,"Interest Level":interest_level,"Links":"",'Path':270}]
            personal_df = pd.concat([personal_df, pd.DataFrame(new_rows)], ignore_index=True)


#print(df.head(70))
personal_df.to_excel(r"C:\Users\mikay\OneDrive\Documents\Resume to Tableau\personal_data.xlsx", index=False)
resume_df.to_excel(r"C:\Users\mikay\OneDrive\Documents\Resume to Tableau\resume_data.xlsx", index=False)
skills_df.to_excel(r'C:\Users\mikay\OneDrive\Documents\Resume to Tableau\skills_data.xlsx', index=False)

import sqlite3
import os

# Connect (creates the file if it doesn't exist)
# Set your desired directory (e.g., Desktop or Documents)
folder_path = os.path.expanduser("~\Documents\Github\Resume-Dashboard")  # or another directory
os.makedirs(folder_path, exist_ok=True)  # Create it if it doesn't exist
# Build the full path to the database
db_path = os.path.join(folder_path, "resume_data.db")
print(db_path)
conn = sqlite3.connect(db_path)

resume_df.to_sql("experience", conn, if_exists="replace", index=False)
personal_df.to_sql('personal',conn, if_exists="replace", index=False)
skills_df.to_sql('skills',conn, if_exists="replace", index=False)

result = pd.read_sql_query("SELECT * FROM experience", conn)
print(result)